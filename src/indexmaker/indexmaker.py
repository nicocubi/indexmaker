import pdfplumber
import argparse
from collections import defaultdict
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.tokenize import word_tokenize
from docx import Document
import nltk
import sys, os

class IndexMaker():

    def __init__(self, pdf_path, output_docx=None):
        self.pdf_path = pdf_path
        if not os.path.isfile(pdf_path):
                    raise FileNotFoundError(f"The file {pdf_path} does not exist.")

        if not output_docx:
            output_docx = os.path.splitext(os.path.basename(pdf_path))[0] + "_index.docx"
        self.output_docx = output_docx if output_docx else "word_index.docx"
        self.pages = {}
        self.index = None

    def process_pdf(self):
        # Extract text page by page
        with pdfplumber.open(self.pdf_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                extract = page.extract_text() or ""
                self.pages[page_number] = extract.lower()

    def create_index(self):
        # Initialize stemmer and stop words
        stemmer = PorterStemmer()
        stop_words = set(stopwords.words("english"))
        word_pages = defaultdict(set)
        stem_to_original = {}

        # Process each page
        for page_number, text in self.pages.items():
            tokens = word_tokenize(text)
            for word in tokens:
                # Skip single-letter words and stop words
                if len(word) > 1 and word.isalpha() and word.lower() not in stop_words:
                    stemmed = stemmer.stem(word.lower())
                    if stemmed not in stem_to_original:
                        stem_to_original[stemmed] = word.lower()
                    word_pages[stem_to_original[stemmed]].add(page_number)

        # Convert to a sorted dictionary
        self.index = {word: sorted(pages) for word, pages in word_pages.items()}

    def save_index_to_docx(self):
        # Create a Word document
        doc = Document()
        doc.add_heading("Word Index", level=1)

        if self.index:
            # Sort words alphabetically
            sorted_words = sorted(self.index.keys())
            current_letter = None

            for word in sorted_words:
                # Group by the first letter
                first_letter = word[0].upper()
                if first_letter != current_letter:
                    current_letter = first_letter
                    doc.add_heading(current_letter, level=2)

                # Add the word and its pages in the same paragraph
                paragraph = doc.add_paragraph(f"{word}: {', '.join(map(str, self.index[word]))}")
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = 0
                paragraph_format.space_after = 0
                paragraph_format.line_spacing = 1  # Optional: Set line spacing to single

            # Save the document
            doc.save(self.output_docx)
        else:
            print("Empty index!")



#----------MAIN

if __name__ == "__main__":


    # Ensure NLTK resources are available
    #nltk.download('punkt')
    #nltk.download('stopwords')
    parser = argparse.ArgumentParser(description="Generate a word index from a PDF file.")
    parser.add_argument("pdf_path", help="Path to the input PDF file.")
    parser.add_argument("--output", help="Path to the output Word document.", default="word_index.docx")
    args = parser.parse_args()
    pdf_path = args.pdf_path

    # Example usage
    index_maker = IndexMaker(pdf_path)
    index_maker.process_pdf()
    index_maker.create_index()
    index_maker.save_index_to_docx()

